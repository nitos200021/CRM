from flask import Flask, render_template, flash, redirect, url_for, make_response, request, send_file
from flask_login import LoginManager, login_required, current_user, login_user
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, TextAreaField, SelectField, SubmitField, DateTimeField
from wtforms.validators import DataRequired
from datetime import datetime
from config import Config
from models import db, User, Employee, Ticket, House, WorkType
from forms import HouseForm, WorkTypeForm, EmployeeForm
from docx import Document
from io import BytesIO

app = Flask(__name__)
app.config.from_object(Config)
db.init_app(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

class LoginForm(FlaskForm):
    username = StringField('Имя пользователя', validators=[DataRequired()])
    password = PasswordField('Пароль', validators=[DataRequired()])
    submit = SubmitField('Войти')

@app.route('/house/new', methods=['GET', 'POST'])
@login_required
def add_house():
    form = HouseForm()
    if form.validate_on_submit():
        new_house = House(
            address=form.address.data,
            floors=form.floors.data
        )
        db.session.add(new_house)
        db.session.commit()
        flash('Адрес успешно добавлен.', 'success')
        return redirect(url_for('index'))
    return render_template('add_house.html', form=form)

@app.route('/work_type/new', methods=['GET', 'POST'])
@login_required
def add_work_type():
    form = WorkTypeForm()
    if form.validate_on_submit():
        new_work_type = WorkType(
            name=form.name.data,
            description=form.description.data
        )
        db.session.add(new_work_type)
        db.session.commit()
        flash('Тип работы успешно добавлен.', 'success')
        return redirect(url_for('index'))
    return render_template('add_work_type.html', form=form)

@app.route('/employee/new', methods=['GET', 'POST'])
@login_required
def add_employee():
    form = EmployeeForm()
    if form.validate_on_submit():
        new_employee = Employee(
            name=form.name.data,
            position=form.position.data
        )
        db.session.add(new_employee)
        db.session.commit()
        flash('Исполнитель успешно добавлен.', 'success')
        return redirect(url_for('index'))
    return render_template('add_employee.html', form=form)

class LoginForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired()])
    password = PasswordField('Password', validators=[DataRequired()])
    submit = SubmitField('Login')

class TicketForm(FlaskForm):
    ticket_number = StringField('Номер заявки', validators=[DataRequired()])
    description = TextAreaField('Описание проблемы', validators=[DataRequired()])
    address = SelectField('Адрес проблемы', coerce=int, validators=[DataRequired()])
    criticality = SelectField('Критичность заявки', choices=[('Низкая','Низкая'),('Средняя','Средняя'),('Высокая','Высокая')], validators=[DataRequired()])
    contact_name = StringField('ФИО', validators=[DataRequired()])
    contact_phone = StringField('Телефон', validators=[DataRequired()])
    work_type = SelectField('Тип работ', coerce=int, validators=[DataRequired()])
    submit = SubmitField('Создать заявку')

class EditTicketForm(FlaskForm):
    description = TextAreaField('Описание проблемы', validators=[DataRequired()])
    status = SelectField('Статус заявки', choices=[('Новая','Новая'),('В работе','В работе'),('Выполнена','Выполнена'),('Отменена','Отменена')], validators=[DataRequired()])
    start_time = DateTimeField('Время начала работы', format='%Y-%m-%dT%H:%M', default=datetime.utcnow)
    end_time = DateTimeField('Время окончания работы', format='%Y-%m-%dT%H:%M')
    assigned_employee = SelectField('Исполнитель', coerce=int, validators=[DataRequired()])
    responsible_person = StringField('Ответственный (мастер смены)', validators=[DataRequired()])
    note = TextAreaField('Примечание')
    submit = SubmitField('Сохранить изменения')

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

def seed_data():
    if not Employee.query.first():
        employees = [Employee(name='Иван Иванов', position='Мастер'), Employee(name='Пётр Петров', position='Специалист'), Employee(name='Сергей Сергеев', position='Инженер')]
        db.session.add_all(employees)
    if not User.query.first():
        user = User(username='admin', password='admin')
        db.session.add(user)
    db.session.commit()

@app.before_first_request
def initialize():
    with app.app_context():
        db.create_all()
        seed_data()

@app.route('/login', methods=['GET','POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user and user.password == form.password.data:
            login_user(user)
            return redirect(url_for('index'))
        flash('Неверный логин или пароль','danger')
    return render_template('login.html', form=form)

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/')
@login_required
def index():
    ticket_query = Ticket.query

    date_filter = request.args.get('date')
    work_filter = request.args.get('work')
    address_filter = request.args.get('address')
    criticality_filter = request.args.get('criticality')
    employee_filter = request.args.get('employee')
    responsible_filter = request.args.get('responsible')
    status_filter = request.args.get('status')
    phone_filter = request.args.get('phone') 

    if date_filter:
        try:
            date_obj = datetime.strptime(date_filter, '%Y-%m-%dT%H:%M')
            ticket_query = ticket_query.filter(db.func.date(Ticket.created_at) == date_obj.date())
        except Exception:
            flash('Неверный формат даты.', 'danger')

    if work_filter:
        ticket_query = ticket_query.join(WorkType).filter(WorkType.id == int(work_filter))

    if address_filter:
        house = House.query.get(int(address_filter))
        if house:
            ticket_query = ticket_query.filter(Ticket.address == house.address)

    if criticality_filter:
        ticket_query = ticket_query.filter(Ticket.criticality == criticality_filter)

    if employee_filter:
        ticket_query = ticket_query.filter(Ticket.assigned_employee_id == int(employee_filter))

    if responsible_filter:
        ticket_query = ticket_query.filter(Ticket.responsible_person == responsible_filter)

    if status_filter:
        ticket_query = ticket_query.filter(Ticket.status == status_filter)

    if phone_filter:
        ticket_query = ticket_query.filter(Ticket.contact_phone.like(f'%{phone_filter}%'))

    tickets = ticket_query.order_by(Ticket.criticality.desc(), Ticket.created_at.desc()).all()

    houses = House.query.all()
    work_types = WorkType.query.all()
    employees = Employee.query.all()

    return render_template(
        'index.html',
        tickets=tickets,
        houses=houses,
        work_types=work_types,
        employees=employees
    )

@app.route('/ticket/new', methods=['GET','POST'])
@login_required
def create_ticket():
    houses = House.query.filter_by(active=True).all()
    work_types = WorkType.query.all()
    form = TicketForm()
    form.address.choices = [(house.id, house.address) for house in houses]
    form.work_type.choices = [(wt.id, wt.name) for wt in work_types]
    if form.validate_on_submit():
        selected_house = House.query.get(form.address.data)
        selected_work = WorkType.query.get(form.work_type.data)
        new_ticket = Ticket(ticket_number=form.ticket_number.data, description=form.description.data, address=selected_house.address, criticality=form.criticality.data, contact_name=form.contact_name.data, contact_phone=form.contact_phone.data, work_type=selected_work)
        db.session.add(new_ticket)
        db.session.commit()
        flash('Заявка успешно создана.','success')
        return redirect(url_for('index'))
    return render_template('create_ticket.html', form=form)

@app.route('/ticket/<int:ticket_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_ticket(ticket_id):
    ticket = Ticket.query.get_or_404(ticket_id)
    employees = Employee.query.all()
    form = EditTicketForm(obj=ticket)
    form.assigned_employee.choices = [(emp.id, f'{emp.name} - {emp.position}') for emp in employees]
    
    if form.validate_on_submit():
        ticket.description = form.description.data
        ticket.status = form.status.data
        ticket.start_time = form.start_time.data
        ticket.end_time = form.end_time.data
        ticket.assigned_employee_id = form.assigned_employee.data
        ticket.responsible_person = form.responsible_person.data
        ticket.note = form.note.data
        db.session.commit()
        flash('Заявка успешно обновлена.', 'success')
        return redirect(url_for('index'))
    
    form.description.data = ticket.description
    form.status.data = ticket.status
    form.start_time.data = ticket.start_time
    form.end_time.data = ticket.end_time
    form.assigned_employee.data = ticket.assigned_employee_id
    form.responsible_person.data = ticket.responsible_person
    form.note.data = ticket.note
    
    return render_template('edit_ticket.html', form=form, ticket=ticket)

@app.route('/act/<int:ticket_id>')
@login_required
def act(ticket_id):
    ticket = Ticket.query.get_or_404(ticket_id)
    house = House.query.filter_by(address=ticket.address).first()
    return render_template('act.html', ticket=ticket, house=house)

@app.route('/act/<int:ticket_id>/download')
@login_required
def act_download(ticket_id):
    ticket = Ticket.query.get_or_404(ticket_id)
    house = House.query.filter_by(address=ticket.address).first()
    
    document = Document()
    document.add_heading(f'Акт выполненных работ № {ticket.ticket_number}', 0)
    document.add_paragraph(f'Дата: {ticket.start_time.strftime("%d.%m.%Y") if ticket.start_time else "__.__.____"}г.')
    document.add_paragraph(f'Адрес: ул. {ticket.address} {"эт. " + house.floors if house and house.floors else ""}')
    document.add_paragraph(f'Контактное лицо: {ticket.contact_name}')
    document.add_paragraph(f'Контактный тел.: {ticket.contact_phone}')
    document.add_paragraph(f'Характер заявки: {ticket.description}')
    document.add_paragraph(f'Выполненные работы: {ticket.work_type.name if ticket.work_type else ""}')
    document.add_paragraph('___________________________________________________')
    document.add_paragraph('___________________________________________________')
    document.add_paragraph('___________________________________________________')
    document.add_paragraph('Материал __________________________________________')
    document.add_paragraph('___________________________________________________')
    document.add_paragraph('___________________________________________________')
    document.add_paragraph(f'Руководитель ___________/_____________/         Исполнитель: {ticket.assigned_employee.name if ticket.assigned_employee else ""}/____________/')
    document.add_paragraph(f'Претензий и замечаний по выполненным работам не имею ____________/____________/')

    f = BytesIO()
    document.save(f)
    f.seek(0)
    return send_file(f, as_attachment=True, attachment_filename=f'act_{ticket.ticket_number}.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/kanban')
@login_required
def kanban():
    statuses = ["Новая", "В работе", "Выполнена", "Отменена"]
    kanban_tickets = {status: Ticket.query.filter_by(status=status).order_by(Ticket.created_at.desc()).all() for status in statuses}
    return render_template("kanban.html", kanban_tickets=kanban_tickets)

@app.route('/reset', methods=['GET','POST'])
@login_required
def reset():
    if request.method == 'POST':
        db.session.query(Ticket).delete()
        db.session.commit()
        flash('Все заявки успешно удалены.','warning')
        return redirect(url_for('index'))
    return render_template('reset.html')

@app.route('/admin/import_addresses', methods=['GET','POST'])
@login_required
def import_addresses():
    if request.method == 'POST':
        file = request.files.get('file')
        if file:
            df = pd.read_excel(file, usecols=["Адрес", "Эт-ть"])
            for index, row in df.iterrows():
                address = row["Адрес"]
                floor = str(row["Эт-ть"])
                if not House.query.filter_by(address=address).first():
                    new_house = House(address=address, floors=floor)
                    db.session.add(new_house)
            db.session.commit()
            flash("Адреса импортированы", "success")
            return redirect(url_for('index'))
    return render_template("import_addresses.html")

@app.route('/admin/import_faults', methods=['GET','POST'])
@login_required
def import_faults():
    if request.method == 'POST':
        file = request.files.get('file')
        if file:
            df = pd.read_excel(file, usecols=["№", "Вид неисправности"])
            for index, row in df.iterrows():
                fault = row["Вид неисправности"]
                if not WorkType.query.filter_by(name=fault).first():
                    new_fault = WorkType(name=fault, description="")
                    db.session.add(new_fault)
            db.session.commit()
            flash("Виды неисправностей импортированы", "success")
            return redirect(url_for('index'))
    return render_template("import_faults.html")

@app.route('/admin/add_address', methods=['GET','POST'])
@login_required
def add_address():
    if request.method == 'POST':
        address = request.form.get('address')
        floor = request.form.get('floor')
        if address and not House.query.filter_by(address=address).first():
            new_house = House(address=address, floors=floor)
            db.session.add(new_house)
            db.session.commit()
            flash("Адрес добавлен", "success")
        else:
            flash("Адрес уже существует или не указан", "warning")
        return redirect(url_for('index'))
    return render_template("add_address.html")

@app.route('/admin/add_fault', methods=['GET','POST'])
@login_required
def add_fault():
    if request.method == 'POST':
        fault_type = request.form.get('fault_type')
        if fault_type and not WorkType.query.filter_by(name=fault_type).first():
            new_fault = WorkType(name=fault_type, description="")
            db.session.add(new_fault)
            db.session.commit()
            flash("Вид неисправности добавлен", "success")
        else:
            flash("Вид неисправности уже существует или не указан", "warning")
        return redirect(url_for('index'))
    return render_template("add_fault.html")

if __name__ == '__main__':
    app.run(debug=True)